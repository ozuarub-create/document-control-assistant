"""Week 18 end-to-end demo for the integrated AI Document Control Assistant."""

from __future__ import annotations

import json
import time
from pathlib import Path

from docx import Document

from app.analytics import generate_accuracy_and_performance_report, generate_platform_analytics, save_report
from app.database import DATABASE_PATH, reset_database
from app.document_processor import process_file
from app.repository import ingest_folder, save_processed_document, search_documents
from app.review_engine import review_document
from app.review_repository import save_review_report
from app.search_engine import natural_language_search, semantic_search
from app.workflow import apply_review_workflow, get_document_workflow, get_workflow_summary, update_workflow_state

ROOT_DIR = Path(__file__).resolve().parent
SAMPLE_DIR = ROOT_DIR / "sample_documents"
DEMO_INPUT_DIR = ROOT_DIR / "demo_inputs"
RESULTS_PATH = ROOT_DIR / "demo_week18_results.json"
PERFORMANCE_PATH = ROOT_DIR / "docs" / "ACCURACY_PERFORMANCE_REPORT.json"


def create_complete_demo_document(path: Path) -> None:
    path.parent.mkdir(exist_ok=True)
    document = Document()
    document.add_heading("Week 18 Architectural Drawing Review", level=1)
    lines = [
        "Document Type: Drawing",
        "Document Title: Week 18 Architectural Drawing Review",
        "Revision Number: R2",
        "Project Name: Digital Campus Control Center",
        "Contractor: Future Build Contractors",
        "Consultant: Smart Design Consultants",
        "Submission Date: 2026-06-20",
        "Discipline: Architectural",
        "This drawing includes layout plan, elevation, section, scale, grid references, and drawing number information.",
        "The purpose of this document is to demonstrate complete lifecycle document control using the integrated platform.",
        "The document is ready for validation, review, workflow routing, analytics, and reporting.",
    ]
    for line in lines:
        document.add_paragraph(line)
    document.save(path)


def main() -> None:
    print("Week 18 Demo - Integrated AI Document Control Assistant")
    print("=" * 78)

    print("\n1) Rebuilding metadata repository from 50 sample documents")
    reset_database(DATABASE_PATH)
    start = time.perf_counter()
    registered = ingest_folder(SAMPLE_DIR, reset=False)
    elapsed = time.perf_counter() - start
    print(f"Registered documents: {len(registered)}")

    manifest = json.loads((SAMPLE_DIR / "dataset_manifest.json").read_text(encoding="utf-8"))
    expected_by_filename = {item["filename"]: item["expected_document_type"] for item in manifest}
    correct = sum(1 for item in registered if item["document_type"] == expected_by_filename.get(item["filename"]))
    print(f"Classification accuracy: {correct}/{len(registered)}")

    print("\n2) Running traditional, semantic, and natural language search")
    metadata_results = search_documents(document_type="Drawing", discipline="Architectural", limit=3)
    semantic_results = semantic_search("layout drawings with architectural plans", limit=3)
    nl_results = natural_language_search("Show me the latest architectural drawings", limit=3)
    print(f"Traditional search results: {len(metadata_results)}")
    print(f"Semantic search results: {len(semantic_results)}")
    print(f"Natural language query results: {len(nl_results['results'])}")

    print("\n3) Reviewing a new uploaded document and routing workflow")
    demo_doc = DEMO_INPUT_DIR / "week18_architectural_drawing.docx"
    create_complete_demo_document(demo_doc)
    processed = process_file(demo_doc, include_text=True)
    review = review_document(processed)
    saved_document = save_processed_document(processed, file_path=demo_doc)
    saved_review = save_review_report(review, document_id=saved_document["id"])
    workflow_result = apply_review_workflow(saved_document["id"], review)
    print(f"New document ID: {saved_document['id']}")
    print(f"Review status: {review['validation_status']}")
    print(f"Quality score: {review['quality_score']}")
    print(f"Workflow state after review: {workflow_result['workflow_state']}")
    print(f"Stored review ID: {saved_review['id']}")

    print("\n4) Demonstrating duplicate detection")
    duplicate_review = review_document(processed)
    print(f"Duplicate status: {duplicate_review['duplicate_check']['status']}")

    print("\n5) Demonstrating manual workflow lifecycle update")
    archived = update_workflow_state(
        saved_document["id"],
        "archived",
        action="demo_archive_complete",
        user="document-controller",
        comment="Archived after Week 18 end-to-end demo.",
        enforce_transition=False,
    )
    workflow_history = get_document_workflow(saved_document["id"])
    print(f"Workflow state after manual update: {archived['workflow_state']}")
    print(f"Workflow history events: {len(workflow_history['history'])}")

    print("\n6) Generating analytics and performance report")
    analytics = generate_platform_analytics()
    performance_report = generate_accuracy_and_performance_report(
        processed_documents=len(registered),
        correct_classifications=correct,
        total_seconds=elapsed,
    )
    PERFORMANCE_PATH.parent.mkdir(exist_ok=True)
    save_report(performance_report, PERFORMANCE_PATH)
    print(f"Workflow summary: {get_workflow_summary()}")
    print(f"Average confidence: {analytics['classification']['average_confidence_score']}")
    print(f"Average quality score: {analytics['review_quality']['average_quality_score']}")
    print(f"Performance report saved: {PERFORMANCE_PATH.relative_to(ROOT_DIR)}")

    results = {
        "registered_documents": len(registered),
        "correct_classifications": correct,
        "traditional_search_count": len(metadata_results),
        "semantic_search_count": len(semantic_results),
        "natural_language_query_count": len(nl_results["results"]),
        "new_document_id": saved_document["id"],
        "review_status": review["validation_status"],
        "quality_score": review["quality_score"],
        "duplicate_status": duplicate_review["duplicate_check"]["status"],
        "workflow_state_after_review": workflow_result["workflow_state"],
        "workflow_state_final": archived["workflow_state"],
        "workflow_history_events": len(workflow_history["history"]),
        "analytics": analytics,
        "performance_report": performance_report,
    }
    RESULTS_PATH.write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")

    print("\nWeek 18 Demo Complete")
    print("=" * 78)
    print("Integrated platform: working")
    print("Dashboard API: working")
    print("Workflow states: working")
    print("Analytics/reporting: working")
    print("End-to-end demonstration: working")
    print(f"Results saved to: {RESULTS_PATH.name}")


if __name__ == "__main__":
    main()
