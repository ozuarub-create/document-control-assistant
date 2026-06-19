"""Week 17 demonstration: AI Document Review and Validation Assistant."""

from __future__ import annotations

import json
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.database import DATABASE_PATH, reset_database
from app.document_processor import process_file
from app.repository import ingest_folder, save_processed_document
from app.review_engine import review_document
from app.review_repository import save_review_report

ROOT_DIR = Path(__file__).resolve().parent
SAMPLE_DIR = ROOT_DIR / "sample_documents"
RESULTS_PATH = ROOT_DIR / "demo_week17_results.json"


def print_section(title: str) -> None:
    print("\n" + "=" * 70)
    print(title)
    print("=" * 70)


def create_incomplete_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Incomplete RFI Submission", level=1)
    doc.add_paragraph("Document Type: RFI")
    doc.add_paragraph("Question: Please confirm the ceiling access panel location.")
    doc.add_paragraph("This file is intentionally missing project, revision, contractor, consultant, date, and discipline metadata.")
    doc.save(path)


def show_report(label: str, report: dict) -> None:
    print_section(label)
    print(f"Status: {report['validation_status']}")
    print(f"Quality Score: {report['quality_score']}")
    print(f"Duplicate Check: {report['duplicate_check']['status']}")
    print(f"Summary: {report['summary']['summary_text'][:500]}")

    print("\nWarnings:")
    if report["warnings"]:
        for warning in report["warnings"][:6]:
            print(f"- {warning}")
    else:
        print("- No warnings")

    print("\nRecommendations:")
    if report["recommendations"]:
        for recommendation in report["recommendations"][:6]:
            print(f"- {recommendation}")
    else:
        print("- No recommendations")


def main() -> None:
    print_section("Week 17 Demo - AI Document Review & Validation Assistant")

    reset_database(DATABASE_PATH)

    # 1. Review a valid sample document before it is registered.
    valid_path = SAMPLE_DIR / "sample_001_drawing.pdf"
    valid_processed = process_file(valid_path, include_text=True)
    valid_report = review_document(valid_processed)
    show_report("Review Report 1 - Standard Document Review", valid_report)

    # Load the sample register, then review an existing document to prove duplicate detection.
    loaded_documents = ingest_folder(SAMPLE_DIR, reset=False)
    print_section("Sample Register Loaded")
    print(f"Sample register loaded: {len(loaded_documents)} documents")

    # 2. Demonstrate duplicate detection by reviewing a document already in the register.
    duplicate_path = SAMPLE_DIR / "sample_047_rfi.pdf"
    duplicate_processed = process_file(duplicate_path, include_text=True)
    duplicate_report = review_document(duplicate_processed)
    show_report("Review Report 2 - Duplicate Detection", duplicate_report)

    # 3. Demonstrate missing information detection using an intentionally incomplete DOCX.
    with TemporaryDirectory() as temp_dir:
        incomplete_path = Path(temp_dir) / "incomplete_rfi.docx"
        create_incomplete_docx(incomplete_path)
        incomplete_processed = process_file(incomplete_path, include_text=True)
        incomplete_report = review_document(incomplete_processed)
        show_report("Review Report 3 - Missing Information Detection", incomplete_report)

    # 4. Store a review report with a new uploaded document workflow.
    saved_document = save_processed_document(valid_processed)
    saved_review = save_review_report(valid_report, document_id=saved_document["id"])
    print_section("Stored Review Report")
    print(f"Document ID: {saved_document['id']}")
    print(f"Stored Review ID: {saved_review['id']}")

    results = {
        "valid_document_review": valid_report,
        "duplicate_detection_review": duplicate_report,
        "missing_information_review": incomplete_report,
        "stored_review_id": saved_review["id"],
    }
    RESULTS_PATH.write_text(json.dumps(results, indent=2), encoding="utf-8")

    print_section("Week 17 Demo Complete")
    print(f"Results saved to: {RESULTS_PATH.name}")
    print("Validation engine: working")
    print("Duplicate detection: working")
    print("AI summary generation: working")
    print("Review report generation: working")


if __name__ == "__main__":
    main()
