from pathlib import Path

from docx import Document

from app.database import reset_database
from app.document_processor import process_file
from app.repository import save_processed_document
from app.review_engine import review_document
from app.review_repository import save_review_report, get_latest_review_for_document

ROOT_DIR = Path(__file__).resolve().parent.parent
SAMPLE_DIR = ROOT_DIR / "sample_documents"


def test_week17_review_report_for_sample_document(tmp_path):
    db_path = tmp_path / "review_test.db"
    reset_database(db_path)

    processed = process_file(SAMPLE_DIR / "sample_001_drawing.pdf", include_text=True)
    report = review_document(processed, db_path=db_path)

    assert report["filename"] == "sample_001_drawing.pdf"
    assert report["document_type"] == "Drawing"
    assert report["summary"]["summary_text"]
    assert isinstance(report["warnings"], list)
    assert isinstance(report["recommendations"], list)
    assert "duplicate_check" in report
    assert report["quality_score"] >= 0


def test_week17_missing_information_is_detected(tmp_path):
    db_path = tmp_path / "missing_info.db"
    reset_database(db_path)

    incomplete_path = tmp_path / "incomplete_rfi.docx"
    doc = Document()
    doc.add_heading("Incomplete RFI", level=1)
    doc.add_paragraph("Document Type: RFI")
    doc.add_paragraph("Question: Please confirm the ceiling access panel location.")
    doc.save(incomplete_path)

    processed = process_file(incomplete_path, include_text=True)
    report = review_document(processed, db_path=db_path)
    missing_fields = {item["field"] for item in report["missing_information"]}

    assert "project_name" in missing_fields
    assert "contractor" in missing_fields
    assert report["validation_status"] in {"Needs Review", "Fail"}


def test_week17_duplicate_detection_finds_existing_document(tmp_path):
    db_path = tmp_path / "duplicate_test.db"
    reset_database(db_path)

    processed = process_file(SAMPLE_DIR / "sample_047_rfi.pdf", include_text=True)
    save_processed_document(processed, db_path=db_path)

    report = review_document(processed, db_path=db_path)

    assert report["duplicate_check"]["status"] != "No Duplicate Found"
    assert len(report["duplicate_check"]["candidates"]) >= 1


def test_week17_review_report_can_be_saved(tmp_path):
    db_path = tmp_path / "saved_review.db"
    reset_database(db_path)

    processed = process_file(SAMPLE_DIR / "sample_002_drawing.docx", include_text=True)
    saved_document = save_processed_document(processed, db_path=db_path)
    report = review_document(processed, db_path=db_path)
    saved_report = save_review_report(report, document_id=saved_document["id"], db_path=db_path)
    loaded_report = get_latest_review_for_document(saved_document["id"], db_path=db_path)

    assert saved_report["id"] == loaded_report["id"]
    assert loaded_report["report"]["filename"] == "sample_002_drawing.docx"
