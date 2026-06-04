"""Document ingestion and processing helpers."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from app.classifier import classify_document
from app.metadata_extractor import extract_metadata

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text_from_pdf(file_path: str | Path) -> str:
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_text_from_docx(file_path: str | Path) -> str:
    doc = Document(str(file_path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_text(file_path: str | Path, filename: str | None = None) -> str:
    path = Path(file_path)
    suffix = Path(filename or path.name).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)

    raise ValueError(f"Unsupported file type: {suffix}. Only PDF and DOCX are supported.")


def process_file(file_path: str | Path, filename: str | None = None) -> dict[str, Any]:
    """Run ingestion, classification, metadata extraction, and confidence scoring."""
    path = Path(file_path)
    original_name = filename or path.name
    suffix = Path(original_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Only .pdf and .docx files are supported.")

    text = extract_text(path, original_name)
    classification = classify_document(text, original_name)
    metadata = extract_metadata(text, original_name)

    return {
        "filename": original_name,
        "file_type": suffix.replace(".", ""),
        "classification": classification.to_dict(),
        "metadata": metadata,
        "text_preview": text[:500],
    }
